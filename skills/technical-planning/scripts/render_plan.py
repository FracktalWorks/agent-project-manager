#!/usr/bin/env python3
"""
render_plan.py — Convert a project_plan.md to DOCX and PDF.

Outputs are written alongside the input MD file:
  project_plan.md   (unchanged — source of truth)
  project_plan.docx
  project_plan.pdf

Usage:
  python .github/skills/technical-planning/scripts/render_plan.py \\
    --input outputs/penrose-v2/project-plan/project_plan.md

  # DOCX only
  python .github/skills/technical-planning/scripts/render_plan.py \\
    --input outputs/.../project_plan.md --format docx

  # PDF only
  python .github/skills/technical-planning/scripts/render_plan.py \\
    --input outputs/.../project_plan.md --format pdf

Dependencies (auto-installed if missing):
  pip install python-docx markdown weasyprint
"""

from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[4]

# Fracktal brand colour (dark grey / charcoal)
BRAND_COLOR = "1A1A2E"
ACCENT_COLOR = "E94560"
TABLE_HEADER_BG = "2E2E4A"


# ---------------------------------------------------------------------------
# Dependency check / auto-install
# ---------------------------------------------------------------------------

def ensure_package(import_name: str, pip_name: str | None = None) -> bool:
    """Try to import a package; install via pip if missing. Returns True if available."""
    try:
        __import__(import_name)
        return True
    except ImportError:
        pkg = pip_name or import_name
        print(f"  Installing {pkg}...")
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", pkg, "--quiet"],
            capture_output=True,
        )
        if result.returncode != 0:
            print(f"  WARNING: Could not install {pkg}. {result.stderr.decode()[:200]}")
            return False
        try:
            __import__(import_name)
            return True
        except ImportError:
            return False


# ---------------------------------------------------------------------------
# DOCX renderer (python-docx)
# ---------------------------------------------------------------------------

def md_to_docx(md_path: Path, docx_path: Path) -> bool:
    """Convert Markdown to DOCX using python-docx with basic MD parsing."""
    if not ensure_package("docx", "python-docx"):
        print("  SKIP DOCX: python-docx not available.")
        return False

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Style helpers ──
    def set_heading_style(para, level: int) -> None:
        run = para.runs[0] if para.runs else para.add_run(para.text)
        sizes = {1: 20, 2: 16, 3: 13, 4: 12}
        run.font.size = Pt(sizes.get(level, 11))
        run.font.bold = True
        r, g, b = (26, 26, 46) if level <= 2 else (46, 46, 74)
        run.font.color.rgb = RGBColor(r, g, b)
        para.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        para.paragraph_format.space_after = Pt(4)

    def add_table_from_md(doc: Document, header: list[str], rows: list[list[str]]) -> None:
        """Render a markdown table into a docx table."""
        col_count = len(header)
        table = doc.add_table(rows=1, cols=col_count)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            cell = hdr_cells[i]
            cell.text = h.strip()
            run = cell.paragraphs[0].runs
            if run:
                run[0].font.bold = True
                run[0].font.size = Pt(9)
                run[0].font.color.rgb = RGBColor(255, 255, 255)
            # Set background colour
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), TABLE_HEADER_BG)
            tcPr.append(shd)

        # Data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                if i < col_count:
                    row_cells[i].text = cell_text.strip()
                    for run in row_cells[i].paragraphs[0].runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

    # ── Parse and render MD ──
    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Heading
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            text = _strip_inline(m.group(2))
            if level == 1:
                para = doc.add_heading(text, level=0)
                para.runs[0].font.size = Pt(20)
                para.runs[0].font.color.rgb = RGBColor(26, 26, 46)
            else:
                para = doc.add_heading(text, level=min(level, 4))
                set_heading_style(para, level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            doc.add_paragraph("_" * 60).paragraph_format.space_before = Pt(4)
            i += 1
            continue

        # Table — collect all rows
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # Filter out separator rows (---|---)
            data_lines = [l for l in table_lines if not re.match(r"^\|[\s\-|:]+\|$", l)]
            if data_lines:
                parsed = [
                    [c.strip() for c in row.strip("|").split("|")]
                    for row in data_lines
                ]
                header = parsed[0] if parsed else []
                rows = parsed[1:] if len(parsed) > 1 else []
                add_table_from_md(doc, header, rows)
            continue

        # Blockquote
        if stripped.startswith("> "):
            text = _strip_inline(stripped[2:])
            para = doc.add_paragraph(text)
            para.paragraph_format.left_indent = Cm(0.8)
            para.runs[0].font.italic = True if para.runs else None
            i += 1
            continue

        # Bullet / list item
        if re.match(r"^[\-\*] ", stripped) or re.match(r"^- \[[ x]\] ", stripped):
            text = re.sub(r"^[\-\*] ", "", stripped)
            text = re.sub(r"^- \[[ x]\] ", "", text)
            text = _strip_inline(text)
            para = doc.add_paragraph(text, style="List Bullet")
            para.runs[0].font.size = Pt(10) if para.runs else None
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        text = _strip_inline(stripped)
        if text:
            para = doc.add_paragraph()
            # Handle **bold** chunks inline
            _add_rich_run(para, stripped)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
        i += 1

    doc.save(str(docx_path))
    print(f"  Saved: {docx_path.relative_to(REPO_ROOT)}")
    return True


def _strip_inline(text: str) -> str:
    """Remove markdown inline markers to get plain text."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
    return text


def _add_rich_run(para, text: str) -> None:
    """Add paragraph runs that preserve **bold** inline formatting."""
    from docx.shared import Pt, RGBColor
    # Split on **...**
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        m = re.match(r"\*\*(.*?)\*\*", part)
        if m:
            run = para.add_run(_strip_inline(m.group(1)))
            run.bold = True
            run.font.size = Pt(10)
        else:
            clean = _strip_inline(part)
            if clean:
                run = para.add_run(clean)
                run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# PDF renderer — tries xhtml2pdf (pure Python), then pypandoc, then HTML fallback
# ---------------------------------------------------------------------------

def _build_html(md_path: Path) -> str:
    """Convert MD to styled HTML string."""
    import markdown as md_lib

    raw_md = md_path.read_text(encoding="utf-8")
    html_body = md_lib.markdown(raw_md, extensions=["tables", "fenced_code", "toc"])

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>
  @page {{ margin: 2cm 2.5cm; size: A4; }}
  body {{ font-family: Arial, sans-serif; font-size: 10pt; line-height: 1.5; color: #1A1A2E; }}
  h1 {{ font-size: 18pt; color: #1A1A2E; border-bottom: 2px solid #E94560; padding-bottom: 4px; }}
  h2 {{ font-size: 13pt; color: #2E2E4A; border-bottom: 1px solid #ccc; margin-top: 18px; }}
  h3 {{ font-size: 11pt; color: #2E2E4A; margin-top: 12px; }}
  h4 {{ font-size: 10pt; color: #444; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 9pt; }}
  th {{ background-color: #2E2E4A; color: white; padding: 5px 8px; text-align: left; }}
  td {{ border: 1px solid #ccc; padding: 4px 8px; vertical-align: top; }}
  tr:nth-child(even) td {{ background-color: #f5f5f8; }}
  blockquote {{ border-left: 3px solid #E94560; padding-left: 10px; color: #555; font-style: italic; }}
  code {{ background: #f0f0f0; padding: 1px 4px; font-size: 9pt; }}
  hr {{ border: none; border-top: 1px solid #ddd; margin: 14px 0; }}
  ul, ol {{ margin: 6px 0 6px 20px; }}
  li {{ margin-bottom: 3px; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""


def md_to_pdf(md_path: Path, pdf_path: Path) -> bool:
    """Convert Markdown to PDF. Tries xhtml2pdf → pypandoc → HTML fallback."""
    if not ensure_package("markdown"):
        print("  SKIP PDF: markdown not available.")
        return False

    html = _build_html(md_path)

    # ── Strategy 1: xhtml2pdf (pure Python, no native deps) ──
    if ensure_package("xhtml2pdf"):
        try:
            from xhtml2pdf import pisa
            with open(str(pdf_path), "wb") as pdf_file:
                result = pisa.CreatePDF(html, dest=pdf_file)
            if not result.err:
                print(f"  Saved: {pdf_path.relative_to(REPO_ROOT)}")
                return True
            else:
                print(f"  WARNING: xhtml2pdf reported errors (err={result.err}), trying next method.")
        except Exception as e:
            print(f"  WARNING: xhtml2pdf failed: {e}")

    # ── Strategy 2: pypandoc (requires pandoc to be installed) ──
    if ensure_package("pypandoc"):
        try:
            import pypandoc
            pypandoc.convert_file(
                str(md_path), "pdf",
                outputfile=str(pdf_path),
                extra_args=["--pdf-engine=pdflatex", "-V", "geometry:margin=2cm"],
            )
            print(f"  Saved: {pdf_path.relative_to(REPO_ROOT)}")
            return True
        except Exception as e:
            print(f"  WARNING: pypandoc/pandoc failed: {e}")

    # ── Strategy 3: HTML fallback ──
    html_path = pdf_path.with_suffix(".html")
    html_path.write_text(html, encoding="utf-8")
    print(f"  PDF not available — saved HTML: {html_path.relative_to(REPO_ROOT)}")
    print("  To get a PDF: open the HTML in Chrome and use File > Print > Save as PDF.")
    return False


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Render project_plan.md → DOCX + PDF")
    p.add_argument("--input", required=True, help="Path to project_plan.md")
    p.add_argument(
        "--format",
        choices=["all", "docx", "pdf"],
        default="all",
        help="Which formats to render (default: all)",
    )
    return p.parse_args()


def main() -> None:
    args = parse_args()
    md_path = Path(args.input).resolve()

    if not md_path.exists():
        print(f"ERROR: File not found: {md_path}", file=sys.stderr)
        sys.exit(1)

    out_dir = md_path.parent
    stem = md_path.stem  # "project_plan"

    print(f"Rendering: {md_path.relative_to(REPO_ROOT)}")

    if args.format in ("all", "docx"):
        docx_path = out_dir / f"{stem}.docx"
        md_to_docx(md_path, docx_path)

    if args.format in ("all", "pdf"):
        pdf_path = out_dir / f"{stem}.pdf"
        md_to_pdf(md_path, pdf_path)

    print("\nDone.")


if __name__ == "__main__":
    main()
